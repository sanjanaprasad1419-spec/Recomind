import os
import io
import re
import logging
import numpy as np
from PIL import Image


logger = logging.getLogger(__name__)

# Configure Tesseract binary path if installed on Windows
TESSERACT_PATHS = [
    r'C:\Program Files\Tesseract-OCR\tesseract.exe',
    r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
    os.path.expanduser(r'~\AppData\Local\Programs\Tesseract-OCR\tesseract.exe')
]

for p in TESSERACT_PATHS:
    if os.path.exists(p):
        try:
            import pytesseract
            pytesseract.pytesseract.tesseract_cmd = p
            logger.info(f"Configured Tesseract binary at: {p}")
            break
        except Exception:
            pass


_EASYOCR_READER_CACHE = None


def get_easyocr_reader():
    """
    Singleton loader for local EasyOCR deep-learning reader (runs on CPU/GPU).
    """
    global _EASYOCR_READER_CACHE
    if _EASYOCR_READER_CACHE is not None:
        return _EASYOCR_READER_CACHE

    try:
        import easyocr
        logger.info("Initializing local EasyOCR Reader...")
        _EASYOCR_READER_CACHE = easyocr.Reader(['en'], gpu=False, verbose=False)
    except Exception as e:
        logger.warning(f"EasyOCR initialization error: {e}")

    return _EASYOCR_READER_CACHE


def extract_text_with_easyocr(image_input) -> str:
    """
    Extracts text locally using EasyOCR deep learning model without external exe requirements.
    """
    try:
        reader = get_easyocr_reader()
        if reader is None:
            return ""
        results = reader.readtext(image_input, detail=0)
        return " ".join(results).strip()
    except Exception as exc:
        logger.warning(f"EasyOCR extraction failed: {exc}")
        return ""


def extract_text_with_gemini_vision(image_bytes: bytes, mime_type: str = "image/png") -> str:
    """
    Calls Google Gemini Vision API (gemini-2.5-flash / gemini-1.5-flash)
    to perform high-accuracy OCR on scanned PDF pages or image notes.
    """
    api_key = os.environ.get("GEMINI_API_KEY") or os.environ.get("AI_API_KEY") or os.environ.get("GEMINI_KEY")
    if not api_key:
        return ""

    import base64
    import requests

    b64_data = base64.b64encode(image_bytes).decode('utf-8')
    models_to_try = ["gemini-2.5-flash", "gemini-1.5-flash"]

    for model_name in models_to_try:
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent?key={api_key}"
        headers = {"Content-Type": "application/json"}
        payload = {
            "contents": [{
                "parts": [
                    {"text": "Perform high-accuracy OCR transcription on this study notes page. Transcribe ALL text, formulas, equations, headings, and bullet points exactly as written. Return ONLY the transcribed text."},
                    {"inlineData": {"mimeType": mime_type, "data": b64_data}}
                ]
            }]
        }
        try:
            res = requests.post(url, headers=headers, json=payload, timeout=20)
            if res.status_code == 200:
                candidates = res.json().get("candidates", [])
                if candidates:
                    parts = candidates[0].get("content", {}).get("parts", [])
                    if parts:
                        return parts[0].get("text", "").strip()
        except Exception as e:
            logger.warning(f"Gemini Vision OCR failed ({model_name}): {e}")

    return ""


def process_note_ocr(file_path: str, file_type: str = "") -> str:
    """
    Multi-Engine OCR & Document Text Extraction Service:
    Extracts text from uploaded study notes (PDF, DOCX, JPG, JPEG, PNG).
    
    Tiers:
    1. Text-based PDF extraction (PyMuPDF / PyPDF)
    2. AI Multimodal Vision OCR (Gemini 2.5 Flash / 1.5 Flash)
    3. Local Deep-Learning OCR (EasyOCR PyTorch Reader)
    4. Tesseract OCR & Word Docx extraction (python-docx)
    """
    if isinstance(file_path, int):
        from api.models import Note
        try:
            note = Note.objects.get(id=file_path)
            file_path = note.uploaded_file.path
            file_type = note.file_type
        except Exception:
            return ""

    if not file_type and isinstance(file_path, str):
        file_type = os.path.splitext(file_path)[1].lower().replace('.', '')

    ext = file_type.lower().replace('.', '')
    extracted_text = ""

    # 1. Multi-page PDF Extraction (PyMuPDF + Gemini Vision OCR + EasyOCR)
    if ext == 'pdf':
        try:
            import fitz  # PyMuPDF
            import concurrent.futures
            doc = fitz.open(file_path)
            page_texts = []
            scanned_pages = []

            for page_idx, page in enumerate(doc):
                txt = page.get_text()
                if txt and len(txt.strip()) > 0:
                    page_texts.append(txt.strip())
                else:
                    scanned_pages.append(page_idx)

            # If document lacks embedded text, OCR up to first 3 scanned pages concurrently
            if not page_texts and scanned_pages:
                target_pages = scanned_pages[:3]
                
                def _ocr_single_page(p_idx):
                    try:
                        p = doc[p_idx]
                        pix = p.get_pixmap(dpi=90)
                        img_bytes = pix.tobytes("png")
                        
                        ocr_txt = extract_text_with_gemini_vision(img_bytes, "image/png")
                        if not ocr_txt or len(ocr_txt.strip()) < 5:
                            img_pil = Image.open(io.BytesIO(img_bytes))
                            img_pil.thumbnail((700, 700))
                            ocr_txt = extract_text_with_easyocr(np.array(img_pil))

                        if not ocr_txt or len(ocr_txt.strip()) < 5:
                            try:
                                import pytesseract
                                img = Image.open(io.BytesIO(img_bytes))
                                ocr_txt = pytesseract.image_to_string(img)
                            except Exception:
                                ocr_txt = ""

                        return ocr_txt.strip() if ocr_txt else ""
                    except Exception as page_ocr_err:
                        logger.warning(f"PDF Page {p_idx} OCR error: {page_ocr_err}")
                        return ""

                with concurrent.futures.ThreadPoolExecutor(max_workers=min(3, os.cpu_count() or 1)) as executor:
                    ocr_results = list(executor.map(_ocr_single_page, target_pages))

                page_texts = [r for r in ocr_results if r and len(r) >= 5]

            extracted_text = "\n\n".join(page_texts)
        except Exception as e:
            logger.warning(f"PyMuPDF note extraction error for {file_path}: {e}. Falling back to pypdf.")
            try:
                from pypdf import PdfReader
                reader = PdfReader(file_path)
                page_texts = [p.extract_text().strip() for p in reader.pages if p.extract_text() and len(p.extract_text().strip()) >= 5]
                extracted_text = "\n\n".join(page_texts)
            except Exception as pypdf_err:
                logger.error(f"PyPDF fallback error for {file_path}: {pypdf_err}")

    # 2. DOCX Word Document Extraction
    elif ext == 'docx':
        try:
            import docx
            doc = docx.Document(file_path)
            full_text = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_txt = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_txt:
                        full_text.append(row_txt)
            extracted_text = "\n".join(full_text)
        except Exception as e:
            logger.error(f"DOCX note extraction error for {file_path}: {e}")

    # 3. Image OCR (JPG, JPEG, PNG)
    elif ext in ['jpg', 'jpeg', 'png']:
        try:
            mime_map = {'jpg': 'image/jpeg', 'jpeg': 'image/jpeg', 'png': 'image/png'}
            with open(file_path, 'rb') as f:
                img_bytes = f.read()
            
            # Primary AI Vision OCR
            extracted_text = extract_text_with_gemini_vision(img_bytes, mime_map.get(ext, 'image/png'))
            
            # Local EasyOCR Fallback
            if not extracted_text or len(extracted_text.strip()) < 5:
                extracted_text = extract_text_with_easyocr(file_path)

            # Tesseract Fallback
            if not extracted_text or len(extracted_text.strip()) < 5:
                import pytesseract
                img = Image.open(file_path)
                extracted_text = pytesseract.image_to_string(img)
        except Exception as e:
            logger.warning(f"Image OCR extraction error for {file_path}: {e}")

    # 4. Direct text read for txt or markdown files
    if not extracted_text or not extracted_text.strip():
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                raw = f.read()
                if len(raw.strip()) >= 5 and not raw.startswith('%PDF'):
                    extracted_text = raw.strip()
        except Exception:
            pass

    return extracted_text.strip()
