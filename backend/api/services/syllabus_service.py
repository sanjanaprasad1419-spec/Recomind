import os
import re
import logging
from PIL import Image

logger = logging.getLogger(__name__)

def clean_text_formatting(text: str) -> str:
    if not text:
        return ""
    text = text.replace('\xa0', ' ').replace('\ufffd', '-').replace('\u2013', '-').replace('\u2014', '-').replace('\u2019', "'")
    text = re.sub(r'[\-\s]+\:\s*', ': ', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def extract_syllabus_text(file_path: str, file_type: str) -> str:
    """
    Extracts text content from an uploaded syllabus file (PDF, DOCX, JPG, JPEG, PNG).
    Preserves all pages without arbitrary truncation.
    """
    ext = file_type.lower().replace('.', '')
    extracted_text = ""

    if ext == 'pdf':
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            page_texts = []
            for page_idx, page in enumerate(reader.pages):
                txt = page.extract_text()
                if txt and txt.strip():
                    page_texts.append(txt.strip())
            extracted_text = "\n\n".join(page_texts)
        except Exception as e:
            logger.error(f"PDF extraction error for {file_path}: {e}")

    elif ext == 'docx':
        try:
            import docx
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_txt = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_txt:
                        full_text.append(row_txt)
            extracted_text = "\n".join(full_text)
        except Exception as e:
            logger.error(f"DOCX extraction error for {file_path}: {e}")

    elif ext in ['jpg', 'jpeg', 'png']:
        try:
            import pytesseract
            img = Image.open(file_path)
            extracted_text = pytesseract.image_to_string(img)
        except Exception as e:
            logger.warning(f"Image OCR extraction fallback for {file_path}: {e}")
            extracted_text = f"Syllabus Image Document: {os.path.basename(file_path)}"

    if not extracted_text or not extracted_text.strip():
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                extracted_text = f.read()
        except Exception:
            extracted_text = f"Syllabus Document: {os.path.basename(file_path)}"

    return extracted_text.strip()


def normalize_concept_lines(raw_lines: list) -> list:
    """
    Normalizes concept lines extracted from syllabus tables/text.
    Combines broken lines into coherent, meaningful educational topic phrases.
    """
    cleaned_topics = []
    current_phrase = []

    for line in raw_lines:
        line = clean_text_formatting(line)
        if not line:
            continue

        cl = re.sub(r'^(?:[•\*\-\+\s]+|\d+(?:\.\d+)*[\.\)\s]*)', '', line).strip()

        is_continuation = bool(re.match(r'^(and|or|of|in|to|the|with|for|a|an|[a-z])\b', cl))
        is_new_concept = (
            (line.startswith('•') or line.startswith('-') or re.match(r'^\d+[\.\)]', line) or re.match(r'^c\d+\.\d+', line, re.IGNORECASE))
            and not is_continuation
        )

        if is_new_concept:
            if current_phrase:
                full_p = " ".join(current_phrase).strip()
                full_p = re.sub(r'\s+', ' ', full_p)
                if len(full_p.split()) >= 1 and len(full_p) >= 3:
                    cleaned_topics.append(full_p)
                current_phrase = []
            if cl:
                current_phrase.append(cl)
        else:
            if current_phrase:
                current_phrase.append(cl)
            else:
                if cl:
                    current_phrase.append(cl)

    if current_phrase:
        full_p = " ".join(current_phrase).strip()
        full_p = re.sub(r'\s+', ' ', full_p)
        if len(full_p.split()) >= 1 and len(full_p) >= 3:
            cleaned_topics.append(full_p)

    unique_topics = []
    seen = set()
    for t in cleaned_topics:
        t_key = t.lower()
        if t_key not in seen and len(t) >= 3 and not re.match(r'^(page|\d+|class|subject code)', t, re.IGNORECASE):
            seen.add(t_key)
            unique_topics.append(t)

    return unique_topics


def parse_syllabus_into_units(syllabus_text: str) -> list:
    """
    Hierarchical Syllabus Parser:
    Extracts Part/Section -> Chapter/Theme -> Syllabus Concepts/Topics.
    
    Handles multi-page syllabi, tables, multi-part syllabi (Part 1, Part 2),
    explicit chapter tags (Chapter 1, CHAPTER III), numbered themes (1. Oceans),
    and Course Outlines without creating broken fragments or fake chapters.
    """
    if not syllabus_text or not isinstance(syllabus_text, str):
        return []

    lines = [clean_text_formatting(line) for line in syllabus_text.splitlines() if line.strip()]

    # Locate Course Outline or main content section if available
    start_idx = 0
    for idx, line in enumerate(lines):
        if 'course outline' in line.lower():
            start_idx = idx
            break

    relevant_lines = lines[start_idx:] if start_idx > 0 else lines

    sections = []
    current_part = "Part 1"
    raw_topic_buffers = {}

    part_pattern = re.compile(
        r'^\s*(part\s*[\d\wIVX]+|section\s*[\d\wA-Z]+|unit\s*[\d\wIVX]+)\b',
        re.IGNORECASE
    )

    explicit_chapter_pattern = re.compile(
        r'^\s*\b(?:chapter|ch|theme|module|unit)\b[\s\-:#]*(\d+|[IVXLCDM]+)\s*(?:[:\-\.\u2013\u2014])?\s*(.+)?$',
        re.IGNORECASE
    )

    numbered_chapter_pattern = re.compile(
        r'^\s*(\d{1,2})(?:[\.\)\-\:\u2013\u2014]|\s+)(?!\d)\s*([A-Za-z][A-Za-z0-9\s,\-\'\(\)]+)$'
    )



    learning_outcome_keywords = [
        'learning outcomes', 'pertinent', 'cgs, cs', 'students will be able to:',
        'explain how', 'describe how', 'understand the', 'analyse how', 'identify',
        'categorise', 'differentiate', 'appreciate', 'explore the', 'recognise',
        'illustrate', 'construct', 'use of', 's. no.', 'theme (time', 'outline/concepts',
        'hours)', 'unit-i', 'unit-ii'
    ]

    i = 0
    N = len(relevant_lines)

    while i < N:
        line = relevant_lines[i]
        line_lower = line.lower()

        # Check for Part/Section Header
        if part_pattern.match(line) and len(line) < 35:
            current_part = line.strip()
            i += 1
            continue

        # Match Chapter or Numbered Theme
        match_exp = explicit_chapter_pattern.match(line)
        match_num = numbered_chapter_pattern.match(line) if not match_exp else None

        num = None
        rest = None

        if match_exp:
            c_num, c_name = match_exp.groups()
            num = c_num.strip()
            rest = c_name.strip() if c_name else ""
        elif match_num:
            c_num, c_name = match_num.groups()
            num = c_num.strip()
            rest = c_name.strip()

        if num is not None:
            title_parts = [rest] if rest else []
            j = i + 1

            # Only look ahead for title continuation if rest is empty or incomplete
            if not title_parts and j < N:
                next_line = relevant_lines[j]
                next_lower = next_line.lower()
                if not re.match(r'^\d+\.\d+', next_line) and not part_pattern.match(next_line) and not explicit_chapter_pattern.match(next_line):
                    title_parts.append(next_line)
                    j += 1

            raw_title = " ".join(title_parts).strip()
            clean_title = re.sub(r'\(\d+\s*hours?\)', '', raw_title, flags=re.IGNORECASE).strip()
            
            concept_split = re.split(r'[\•\-\:]', clean_title, maxsplit=1)
            theme_name = concept_split[0].strip()
            theme_name = re.sub(r'^\d+[\.\)]\s*', '', theme_name).strip()

            if theme_name and len(theme_name.split()) >= 1 and len(theme_name) >= 3:
                part_slug = re.sub(r'[^a-zA-Z0-9]', '', current_part.lower())
                if not part_slug:
                    part_slug = "part1"

                theme_idx = len([s for s in sections if s["part"] == current_part]) + 1
                unique_id = f"chapter_{num}_{theme_idx}" if "part" not in current_part.lower() else f"{part_slug}-theme{num or theme_idx}"

                is_explicit_ch = "chapter" in line_lower or "ch." in line_lower or "ch " in line_lower
                display_label = f"Chapter {num} \u2014 {theme_name}" if is_explicit_ch else f"{num} \u2014 {theme_name}"

                current_theme = {
                    "id": unique_id,
                    "unit_id": unique_id,
                    "number": str(num),
                    "name": theme_name,
                    "title": display_label,
                    "part": current_part,
                    "topics": []
                }
                sections.append(current_theme)
                raw_topic_buffers[unique_id] = []

                if len(concept_split) > 1 and concept_split[1].strip():
                    raw_topic_buffers[unique_id].append(concept_split[1].strip())

                i = j
                continue



        # Topic/Concept Collection
        if sections:
            curr = sections[-1]
            curr_id = curr["id"]
            is_lo = any(lo_kw in line_lower for lo_kw in learning_outcome_keywords) or re.match(r'^c\d+\.\d+', line_lower)

            
            if not is_lo and len(line) >= 3 and not line.endswith(' Hours)'):
                if not re.match(r'^(page\s+\d+|\d+\s*$|class\s+\d+|subject code)', line, re.IGNORECASE):
                    raw_topic_buffers[curr_id].append(line)


        i += 1

    # Normalize concept lines
    for s in sections:
        s_id = s["id"]
        raw_list = raw_topic_buffers.get(s_id, [])
        s["topics"] = normalize_concept_lines(raw_list)

    # Deduplicate sections
    deduped_sections = []
    seen_ids = set()
    for s in sections:
        if s["id"] not in seen_ids:
            seen_ids.add(s["id"])
            deduped_sections.append(s)

    if not deduped_sections:
        logger.warning("Could not detect structured chapters/themes in uploaded syllabus text.")

    return deduped_sections
